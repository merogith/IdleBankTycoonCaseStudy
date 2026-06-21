#!/usr/bin/env python3
# Generate ATS-safe PDF + DOCX for the CV and cover letter (single column,
# standard Arial-compatible font, selectable text, no tables/graphics).
#
# These are the parser-friendly copies to upload into application forms (Lever, etc.).
# The designed copies (cv.html / cover-letter.html) are printed to PDF from a browser.
#
# Usage:  python3 scripts/build-ats-docs.py
# Needs:  pip install reportlab python-docx   (Liberation Sans fonts, preinstalled on most Linux)
# Output: Meric_Erler_CV.{pdf,docx}, Meric_Erler_Cover_Letter.{pdf,docx} in the repo root.
import os
from reportlab.lib.pagesizes import A4
from reportlab.lib.units import mm
from reportlab.lib.enums import TA_LEFT
from reportlab.lib import colors
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.lib.styles import ParagraphStyle
from reportlab.platypus import (SimpleDocTemplate, Paragraph, Spacer,
                                ListFlowable, ListItem, HRFlowable)
import docx
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT = os.path.abspath(os.path.join(os.path.dirname(__file__), ".."))

# ---- fonts (PDF) ----
LIB = "/usr/share/fonts/truetype/liberation"
pdfmetrics.registerFont(TTFont("Sans", f"{LIB}/LiberationSans-Regular.ttf"))
pdfmetrics.registerFont(TTFont("Sans-Bold", f"{LIB}/LiberationSans-Bold.ttf"))
pdfmetrics.registerFont(TTFont("Sans-Italic", f"{LIB}/LiberationSans-Italic.ttf"))
pdfmetrics.registerFontFamily("Sans", normal="Sans", bold="Sans-Bold",
                              italic="Sans-Italic")

# ---------- shared content model ----------
# block kinds: h1, subtitle, contact(list), h2, p(segments), bullets(list of segments),
#              job(segments), spacer
NAME = "Meriç Erler"
TITLE = "Game Economy Designer - F2P & Idle Systems"

CONTACT_CV = [
    "Berlin, Germany  |  mericerler@gmail.com  |  linkedin.com/in/mericerler  |  github.com/merogith",
    "EU work-authorized  |  English (full professional), Turkish (native)  |  Available immediately",
]

def S(*pairs):  # build a list of (text, bold) segments
    return list(pairs)

cv_blocks = [
    ("h1", NAME),
    ("subtitle", TITLE),
    ("contact", CONTACT_CV),
    ("h2", "Summary"),
    ("p", S((
        "Game Economy Designer focused on free-to-play (F2P) and idle games, with three years of "
        "data analysis across pharma, edtech and equity research. I balance game economies by the "
        "numbers: sources and sinks, inflation control, progression curves, retention, and "
        "monetization. To demonstrate the skill set directly, I built a full, self-directed "
        "economy-design study on Idle Bank Tycoon - a parameter-driven model in Excel and Google "
        "Sheets (400+ linked formulas) with two live-ops levers, ROI and monetization math, "
        "sensitivity analysis, and an A/B test plan, modelling a blended ARPDAU lift of +15.5%. I "
        "also build and ship my own browser games, so I read an economy from both the spreadsheet "
        "and the code.", False))),
    ("h2", "Core Skills"),
    ("bullets", [
        S(("Game economy & F2P: ", True), ("economy modelling and balancing; sources/sinks and inflation control; progression curves; monetization (IAP, rewarded ads, ARPDAU, ARPPU, eCPM); retention (D1/D7/D30); player psychology, motivation and segmentation.", False)),
        S(("Spreadsheet modelling: ", True), ("advanced Excel and Google Sheets - parameter-driven models, 400+ linked formulas, sensitivity grids, scenario and A/B design.", False)),
        S(("Data analysis & methods: ", True), ("Python (pandas, scikit-learn, statsmodels); SQL (window functions); A/B testing; regression; z-tests; logistic models; forecasting.", False)),
        S(("BI & reporting: ", True), ("Tableau, Power BI, Looker Studio, Google Analytics, KPI dashboards.", False)),
        S(("Build / technical: ", True), ("custom game engines in JavaScript / TypeScript (Canvas, game loop, rules-vs-render separation); concepts that carry over to Unity; Supabase; Git.", False)),
    ]),
    ("h2", "Selected Projects"),
    ("bullets", [
        S(("Idle Bank Tycoon - Game Economy Design Case Study ", True), ("(Excel/Google Sheets model, A/B design). Found the late-game problem in the numbers: by level 60, idle cash accumulates to 4,900-21,600x the next upgrade's cost, so progression stops mattering. Designed two live-ops levers, a hard-currency liquidity lock and income-scaled rewarded-ad tuning, and modelled each to a blended ARPDAU +15.5% (about $2M/year at static DAU), with a sensitivity grid and retention guardrails. My recommendation: ship the ad tuning first and treat the loan as the bigger retention bet.", False)),
        S(("Free-to-Play Game Economy & Balance ", True), ("(Python, scikit-learn, SQL). Caught two premium items breaking the game (61% win rate, significant on a z-test) and slow coin inflation (sinks recover only ~63% of currency minted). Predicted likely future payers from their first three days (AUC 0.80) using D1/D7/D30 cohorts and turned it into an onboarding lever.", False)),
        S(("US Stock-Market Valuation & Risk ", True), ("(Python, statsmodels, SQL). Across 150+ years, higher valuation (CAPE) predicts lower forward 10-year real returns (p ~ 1e-40); worst real drawdown -81%.", False)),
        S(("Playable browser games ", True), ("(custom JavaScript engines, Supabase). Battle, a turn-based RPG with online PvP; RocketIO, a real-time hex-strategy game with 14 factions and an in-game missile economy. Built mobile-first and playable on phone and desktop. Portfolio: merogith.github.io/Portfolio", False)),
        S(("Idle games roadmap. ", True), ("Building toward an idle tycoon and a single-player idle RPG. Idle systems are where I want to spend the next few years.", False)),
    ]),
    ("h2", "Experience"),
    ("job", S(("Independent Financial Coach", True), ("  -  Self-employed  (2024 - Present)", False))),
    ("bullets", [S(("Run a one-to-one practice turning budgeting, saving and investing into plain, actionable decisions; design each client's plan and run weekly check-ins in English and Turkish.", False))]),
    ("job", S(("Omnichannel Data Analyst", True), ("  -  Viatris Inc.  (Dec 2022 - Sep 2023)", False))),
    ("bullets", [S(("Consolidated records into one master record per person across 70+ markets (Salesforce); fuzzy matching cut manual validation by 40%; cleaned 1M+ records with SQL.", False))]),
    ("job", S(("Business Intelligence Analyst", True), ("  -  Saksı Kampüs (early-stage edtech)  (Dec 2021 - Oct 2022)", False))),
    ("bullets", [S(("Tracked retention, engagement and CAC in Google Analytics and Looker Studio and steered Meta Ads spend; wrote the business plan and growth projections behind a seed crowdfunding raise.", False))]),
    ("job", S(("Sales & Financial Analysis Intern", True), ("  -  Garanti BBVA Investment  (Jul 2021 - Sep 2021)", False))),
    ("bullets", [S(("Built DCF and comparable-company valuation models and set target prices for sell-side equity research.", False))]),
    ("h2", "Education"),
    ("bullets", [
        S(("M.Sc. Corporate Management - University of Europe for Applied Sciences, Berlin (2022 - 2024).", False)),
        S(("B.Sc. Environmental Engineering - METU, Ankara (2013 - 2021).", False)),
    ]),
    ("p", S(("Honors: ", True), ('published paper "Generative AI in Pharma" (ADCAIJ); Google AI Essentials; METU summer-practice award.', False))),
    ("h2", "Languages"),
    ("p", S(("English (full professional); Turkish (native); Portuguese (dual citizen).", False))),
]

P1 = ("I take idle games apart for fun. I will be a few prestiges deep in one and catch myself "
      "reverse-engineering the curve instead of playing it: where is the sink, why did this reward "
      "land now, what holds a player who already has everything? That instinct is why I am writing "
      "about the Game Economy Designer role.")
P2 = ("Rather than tell you I can do the work, I built it. I ran a full, self-directed economy-design "
      "study on Idle Bank Tycoon. Working in Excel and Google Sheets, I built a parameter-driven "
      "model (400+ live formulas) that surfaces the late-game problem: idle cash piling up to "
      "thousands of times the next upgrade's cost. Then I designed two live-ops levers for it, a "
      "hard-currency liquidity lock and income-scaled rewarded-ad tuning, each modelled in full with "
      "ROI and monetization math, a sensitivity grid, retention guardrails and an A/B plan. Together "
      "they land a "
      "blended ARPDAU lift of +15.5% (about $2M/year at static DAU). If I had to sequence them, I "
      "would ship the ad tuning first and treat the loan as the bigger retention bet. The full study "
      "sits alongside this letter.")
P3 = ("My route into economy design went through data: three years turning messy datasets into "
      "decisions across pharma, edtech and equity research. Games are where that analytical side and "
      "the design side finally meet for me. I also build and ship my own browser games on engines I "
      "wrote myself, so I am comfortable moving between the model and the code that runs it. And idle "
      "is honestly next on my own roadmap, with an idle tycoon and a single-player idle RPG as the "
      "two projects I most want to build.")
P4 = ("I would love to bring that mix to the Idle Bank Tycoon team: the analytical rigour from years "
      "in data, a real feel for how these systems hold players, and a genuine love of the genre. "
      "Kolibri helped define idle games with Idle Miner Tycoon, and Idle Bank is the one I would "
      "most like to help push to its next milestones. Thanks for reading, and I would welcome the "
      "chance to talk it through.")

cl_blocks = [
    ("h1", NAME),
    ("subtitle", TITLE),
    ("contact", ["Berlin, Germany  |  mericerler@gmail.com  |  linkedin.com/in/mericerler  |  github.com/merogith"]),
    ("spacer", 6),
    ("p", S(("21 June 2026", False))),
    ("p", S(("Kolibri Games - Game Economy Designer", False))),
    ("spacer", 4),
    ("p", S(("Dear Kolibri Games Team,", False))),
    ("p", S((P1, False))),
    ("p", S((P2, False))),
    ("p", S((P3, False))),
    ("p", S((P4, False))),
    ("spacer", 4),
    ("p", S(("Best regards,", False))),
    ("p", S((NAME, False))),
    ("p", S(("Case study: merogith.github.io/IdleBankTycoonCaseStudy  |  Portfolio: merogith.github.io/Portfolio", False))),
]

# ---------------- PDF renderer ----------------
def esc(t):
    return t.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")

def seg_markup(segments):
    out = ""
    for text, bold in segments:
        e = esc(text)
        out += f"<b>{e}</b>" if bold else e
    return out

def build_pdf(path, blocks, body_size=10.5, leading=14):
    doc = SimpleDocTemplate(path, pagesize=A4,
                            leftMargin=15*mm, rightMargin=15*mm,
                            topMargin=10*mm, bottomMargin=10*mm,
                            title=os.path.basename(path), author=NAME)
    st_h1 = ParagraphStyle("h1", fontName="Sans-Bold", fontSize=17, leading=20, spaceAfter=2, textColor=colors.black)
    st_sub = ParagraphStyle("sub", fontName="Sans-Bold", fontSize=10.5, leading=13, spaceAfter=4, textColor=colors.black)
    st_contact = ParagraphStyle("ct", fontName="Sans", fontSize=8.5, leading=11, spaceAfter=1, textColor=colors.black)
    st_h2 = ParagraphStyle("h2", fontName="Sans-Bold", fontSize=11, leading=13, spaceBefore=7, spaceAfter=2, textColor=colors.black)
    st_body = ParagraphStyle("body", fontName="Sans", fontSize=body_size, leading=leading, spaceAfter=4, alignment=TA_LEFT, textColor=colors.black)
    st_job = ParagraphStyle("job", fontName="Sans", fontSize=body_size, leading=leading, spaceBefore=4, spaceAfter=1, textColor=colors.black)
    st_bul = ParagraphStyle("bul", fontName="Sans", fontSize=body_size, leading=leading, textColor=colors.black)
    flow = []
    for b in blocks:
        kind = b[0]
        if kind == "h1":
            flow.append(Paragraph(esc(b[1]), st_h1))
        elif kind == "subtitle":
            flow.append(Paragraph(esc(b[1]), st_sub))
        elif kind == "contact":
            for line in b[1]:
                flow.append(Paragraph(esc(line), st_contact))
        elif kind == "h2":
            flow.append(Paragraph(esc(b[1]), st_h2))
            flow.append(HRFlowable(width="100%", thickness=0.6, color=colors.black, spaceBefore=1, spaceAfter=3))
        elif kind == "p":
            flow.append(Paragraph(seg_markup(b[1]), st_body))
        elif kind == "job":
            flow.append(Paragraph(seg_markup(b[1]), st_job))
        elif kind == "bullets":
            items = [ListItem(Paragraph(seg_markup(s), st_bul), leftIndent=14, value=None) for s in b[1]]
            flow.append(ListFlowable(items, bulletType="bullet", start="•",
                                     bulletFontName="Sans", bulletFontSize=6,
                                     leftIndent=12, spaceBefore=0, spaceAfter=3))
        elif kind == "spacer":
            flow.append(Spacer(1, b[1]))
    doc.build(flow)
    print("wrote", path)

# ---------------- DOCX renderer ----------------
def set_base_font(d, name="Arial", size=11):
    style = d.styles["Normal"]
    style.font.name = name
    style.font.size = Pt(size)
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    for a in ("w:ascii", "w:hAnsi", "w:cs"):
        rfonts.set(qn(a), name)

def add_bottom_border(par):
    p = par._p
    pPr = p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "000000")
    pbdr.append(bottom)
    pPr.append(pbdr)

def add_runs(par, segments, size=11):
    for text, bold in segments:
        r = par.add_run(text)
        r.bold = bold
        r.font.size = Pt(size)
        r.font.name = "Arial"

def build_docx(path, blocks, body_size=11):
    d = docx.Document()
    set_base_font(d, "Arial", body_size)
    for sec in d.sections:
        sec.top_margin = Inches(0.6); sec.bottom_margin = Inches(0.6)
        sec.left_margin = Inches(0.7); sec.right_margin = Inches(0.7)
    for b in blocks:
        kind = b[0]
        if kind == "h1":
            p = d.add_paragraph(); p.paragraph_format.space_after = Pt(2)
            r = p.add_run(b[1]); r.bold = True; r.font.size = Pt(18); r.font.name = "Arial"
        elif kind == "subtitle":
            p = d.add_paragraph(); p.paragraph_format.space_after = Pt(3)
            r = p.add_run(b[1]); r.bold = True; r.font.size = Pt(12); r.font.name = "Arial"
        elif kind == "contact":
            for line in b[1]:
                p = d.add_paragraph(); p.paragraph_format.space_after = Pt(1)
                r = p.add_run(line); r.font.size = Pt(9); r.font.name = "Arial"
        elif kind == "h2":
            p = d.add_paragraph(); p.paragraph_format.space_before = Pt(10); p.paragraph_format.space_after = Pt(4)
            r = p.add_run(b[1]); r.bold = True; r.font.size = Pt(12); r.font.name = "Arial"
            add_bottom_border(p)
        elif kind == "p":
            p = d.add_paragraph(); p.paragraph_format.space_after = Pt(6)
            add_runs(p, b[1], body_size)
        elif kind == "job":
            p = d.add_paragraph(); p.paragraph_format.space_before = Pt(5); p.paragraph_format.space_after = Pt(1)
            add_runs(p, b[1], body_size)
        elif kind == "bullets":
            for s in b[1]:
                p = d.add_paragraph(style="List Bullet"); p.paragraph_format.space_after = Pt(4)
                add_runs(p, s, body_size)
        elif kind == "spacer":
            d.add_paragraph()
    d.save(path)
    print("wrote", path)

# CV slightly tighter to stay near one page
build_pdf(f"{OUT}/Meric_Erler_CV.pdf", cv_blocks, body_size=9.1, leading=11.2)
build_docx(f"{OUT}/Meric_Erler_CV.docx", cv_blocks, body_size=10.5)
build_pdf(f"{OUT}/Meric_Erler_Cover_Letter.pdf", cl_blocks, body_size=11, leading=15)
build_docx(f"{OUT}/Meric_Erler_Cover_Letter.docx", cl_blocks, body_size=11)
print("DONE")
